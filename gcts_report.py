"""
GCTS Lookup -> Word report generator
-------------------------------------
Reads the "All 35 GCTS Lookup" Excel sheet, drops any row whose GCTS
number is highlighted yellow or green, groups the rest by brand
(AEG and Electrolux merged into one table, every other brand gets its
own table), and writes a formatted .docx with an "Expired Standard"
column filled in and merged down each table.

Usage:
    python gcts_report.py INPUT.xlsx [OUTPUT.docx]

Requires: openpyxl, python-docx  (pip install openpyxl python-docx)
"""

import sys
from copy import deepcopy

import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------- config --

# Fill colors (ARGB hex, as stored by openpyxl) that mark a GCTS row to
# exclude. Add more codes here if the workbook uses other highlight colors.
EXCLUDED_FILL_COLORS = {"FFFFFF00", "FF00B050"}  # yellow, green

# Column letters/positions in the source sheet (1-indexed)
COL_GCTS = 1      # A: GCTS
COL_PRODUCT = 8   # D: Product Name
COL_BRAND = 9     # E: Brand
COL_MODEL = 3     # I: Model No.

SHEET_NAME = "All 35 GCTS Lookup"
HEADER_ROW = 1
FIRST_DATA_ROW = 2

MERGED_BRANDS = {"AEG", "Electrolux"}
MERGED_LABEL = "AEG & Electrolux"

EXPIRED_STANDARD_LINE1 = "IEC 60335-2-6:2014+AMD1:2018"
EXPIRED_STANDARD_LINE2 = "Expired on 23 Aug 2026"

HEADER_BLUE = RGBColor(0x1F, 0x4E, 0x79)
GCTS_BLUE = RGBColor(0x1F, 0x4E, 0x79)
STANDARD_OLIVE = RGBColor(0x8C, 0x8C, 0x00)
EXPIRED_RED = RGBColor(0xFF, 0x00, 0x00)
BORDER_COLOR = "8EA9C1"

COL_WIDTHS_CM = [2.3, 3.9, 2.3, 5.6, 2.5]  # GCTS, Model, Brand, Product, Expired


# ------------------------------------------------------------- extraction --

def load_rows(xlsx_path):
    """Return list of dicts for every non-highlighted row in the sheet."""
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb[SHEET_NAME] if SHEET_NAME in wb.sheetnames else wb.active

    rows = []
    for r in range(FIRST_DATA_ROW, ws.max_row + 1):
        gcts_cell = ws.cell(row=r, column=COL_GCTS)
        if gcts_cell.value is None:
            continue

        fill = gcts_cell.fill.fgColor.rgb if gcts_cell.fill and gcts_cell.fill.fgColor else None
        if fill in EXCLUDED_FILL_COLORS:
            continue  # skip highlighted rows

        rows.append({
            "gcts": str(gcts_cell.value).strip(),
            "model": str(ws.cell(row=r, column=COL_MODEL).value or "").strip(),
            "brand": str(ws.cell(row=r, column=COL_BRAND).value or "").strip(),
            "product": str(ws.cell(row=r, column=COL_PRODUCT).value or "").strip(),
        })
    return rows


def group_by_brand(rows):
    """Group rows by brand, merging MERGED_BRANDS into one group.
    Returns an ordered dict: merged group first, then other brands A-Z."""
    groups = {}
    for row in rows:
        key = MERGED_LABEL if row["brand"] in MERGED_BRANDS else row["brand"]
        groups.setdefault(key, []).append(row)

    ordered_keys = []
    if MERGED_LABEL in groups:
        ordered_keys.append(MERGED_LABEL)
    ordered_keys += sorted(k for k in groups if k != MERGED_LABEL)

    return [(k, groups[k]) for k in ordered_keys]


# ------------------------------------------------------------- docx build --

def set_cell_background(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_borders(cell, color=BORDER_COLOR, size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def write_cell(cell, runs, bold=False, color=None, size=10, align_center_v=True):
    """runs: string, or list of (text, bold, color) tuples for multi-line cells."""
    cell.text = ""
    paragraphs_needed = runs if isinstance(runs, list) else [(runs, bold, color)]
    first = True
    for text, r_bold, r_color in paragraphs_needed:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = "Arial"
        run.bold = r_bold
        if r_color:
            run.font.color.rgb = r_color
    if align_center_v:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_borders(cell)


def add_group_table(doc, title, rows):
    doc.add_heading(f"{title} ({len(rows)})", level=2)

    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.autofit = False

    headers = ["GCTS", "Model No.", "Brand", "Product Name", "Expired Standard"]
    for c, text in enumerate(headers):
        cell = table.rows[0].cells[c]
        write_cell(cell, [(text, True, RGBColor(0xFF, 0xFF, 0xFF))], size=10)
        set_cell_background(cell, "1F4E79")

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        write_cell(cells[0], [(row["gcts"], True, GCTS_BLUE)])
        write_cell(cells[1], row["model"])
        write_cell(cells[2], row["brand"])
        write_cell(cells[3], row["product"])
        if i == 1:
            write_cell(cells[4], [
                (EXPIRED_STANDARD_LINE1, False, STANDARD_OLIVE),
                (EXPIRED_STANDARD_LINE2, True, EXPIRED_RED),
            ])
        else:
            write_cell(cells[4], "")

    # merge the Expired Standard column down the whole table
    if len(rows) > 1:
        first_cell = table.rows[1].cells[4]
        last_cell = table.rows[len(rows)].cells[4]
        first_cell.merge(last_cell)

    # column widths
    for col_idx, width_cm in enumerate(COL_WIDTHS_CM):
        for row in table.rows:
            row.cells[col_idx].width = Cm(width_cm)

    doc.add_paragraph()  # spacing after table


def build_report(rows, output_path, note_extra=""):
    doc = Document()

    title = doc.add_heading("GCTS Lookup — All Entries Except Highlighted", level=1)
    for run in title.runs:
        run.font.color.rgb = HEADER_BLUE

    note = doc.add_paragraph()
    note_run = note.add_run(
        f"Note: GCTS numbers highlighted yellow or green in the source sheet are "
        f"excluded ({len(rows)} rows shown). All entries share the same expired "
        f"standard: {EXPIRED_STANDARD_LINE1}, {EXPIRED_STANDARD_LINE2}.{note_extra}"
    )
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for brand_label, brand_rows in group_by_brand(rows):
        add_group_table(doc, brand_label, brand_rows)

    doc.save(output_path)


# ------------------------------------------------------------------ main --

def main():
    if len(sys.argv) >= 2:
        input_path = sys.argv[1]
        output_path = sys.argv[2] if len(sys.argv) > 2 else "GCTS_Report.docx"
    else:
        # No command-line args (e.g. running via VS Code's "Run" button) —
        # ask interactively instead of exiting.
        input_path = input("Path to the source .xlsx file: ").strip().strip('"')
        output_path = input("Output .docx filename [GCTS_Report.docx]: ").strip().strip('"')
        if not output_path:
            output_path = "GCTS_Report.docx"

    rows = load_rows(input_path)
    if not rows:
        print("No non-highlighted rows found — check EXCLUDED_FILL_COLORS / sheet name.")
        sys.exit(1)

    build_report(rows, output_path)
    print(f"Wrote {output_path} ({len(rows)} rows, "
          f"{len(group_by_brand(rows))} tables)")


if __name__ == "__main__":
    main()